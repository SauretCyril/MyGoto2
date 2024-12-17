from docx import Document
import json
import os

# ...existing code...

def extract_word_to_json(file_path, json_path):
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"The file {file_path} does not exist.")
    
    try:
        doc = Document(file_path)
    except Exception as e:
        raise Exception(f"Failed to load the document: {e}")

    data = {
        "paragraphs": [],
        "tables": [],
        "text_boxes": []
    }

    # Extract paragraphs
    try:
        for paragraph in doc.paragraphs:
            data["paragraphs"].append(paragraph.text)
    except Exception as e:
        print(f"Error extracting paragraphs: {e}")

    # Extract tables
    try:
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            data["tables"].append(table_data)
    except Exception as e:
        print(f"Error extracting tables: {e}")

    # Extract text from inline shapes (including text boxes)
    try:
        for shape in doc.inline_shapes:
            if shape.type == 3:  # 3 corresponds to a text box
                text_box = shape._inline.graphic.graphicData.xpath('.//a:t', namespaces=shape._inline.nsmap)
                for text in text_box:
                    data["text_boxes"].append(text.text)
    except Exception as e:
        print(f"Error extracting text boxes from inline shapes: {e}")

    # Extract text from shapes (including text boxes)
    try:
        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        for shape in doc.element.xpath('//w:txbxContent', namespaces=nsmap):
            print(f"Processing shape: {shape}")
            for paragraph in shape.xpath('.//w:p', namespaces=nsmap):
                print(f"Processing paragraph: {paragraph}")
                texts = [node.text for node in paragraph.xpath('.//w:t', namespaces=nsmap)]
                data["text_boxes"].append(''.join(texts))
    except Exception as e:
        print(f"Error extracting text boxes from shapes: {e}")

    try:
        with open(json_path, 'w', encoding='utf-8') as json_file:
            json.dump(data, json_file, ensure_ascii=False, indent=4)
    except Exception as e:
        print(f"Error writing to JSON file: {e}")

# Example usage:
# extract_word_to_json('path/to/your/document.docx', 'path/to/your/output.json')

extract_word_to_json("G:/OneDrive/Entreprendre/CV/CV_nov_2024_model.docx", "G:/OneDrive/Entreprendre/CV/CV_nov_2024_model.json")
