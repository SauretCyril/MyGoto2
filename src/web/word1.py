from docx import Document

def read_docx(file_path):
    doc = Document(file_path)
    all_text = []

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                print(f"Table Cell: {cell.text}")
                all_text.append(cell.text)

    # Extract text from paragraphs
    for para in doc.paragraphs:
        print(f"Paragraph: {para.text}")
        all_text.append(para.text)

    # Extract text from headers
    for section in doc.sections:
        header = section.header
        for para in header.paragraphs:
            print(f"Header Paragraph: {para.text}")
            all_text.append(para.text)

    # Extract text from footers
    for section in doc.sections:
        footer = section.footer
        for para in footer.paragraphs:
            print(f"Footer Paragraph: {para.text}")
            all_text.append(para.text)

    return '\n'.join(all_text)

file_path = "G:/OneDrive/Entreprendre/CV/CV_nov_2024_model.docx"
extracted_text = read_docx(file_path)
print(extracted_text)